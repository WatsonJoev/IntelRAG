"""DOCX parser using python-docx."""
from __future__ import annotations

from docx import Document as DocxDocument
from docx.document import Document as DocxDocumentType

from ingestion.parsers.base import BaseParser, ParsedDocument


class DOCXParser(BaseParser):
    supported_extensions = [".docx"]
    supported_mime_types = [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ]

    def parse(self, content: bytes, filename: str) -> ParsedDocument:
        import io
        doc = DocxDocument(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n\n".join(paragraphs)
        core = getattr(doc, "core_properties", None)
        title = ""
        author = None
        if core:
            title = (getattr(core, "title", None) or "").strip() or filename
            author = getattr(core, "author", None)
        return ParsedDocument(
            text=text.strip(),
            metadata={"author": author},
            page_count=max(1, len(doc.sections)),
            title=title,
        )
